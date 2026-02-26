import eel
import wx
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
from source.stateManager import GetJudgesData

@eel.expose
def GetDirrectoryPath():
    app = wx.App()
    dlg = wx.DirDialog(None, "Выбор директории...", "",
         wx.DD_DEFAULT_STYLE | wx.DD_DIR_MUST_EXIST)
 
    res = dlg.ShowModal()
    if res == wx.ID_OK:
        print("Выбран каталог: "+dlg.GetPath())
        return dlg.GetPath()

@eel.expose
def CreateWordFile(events, beginDate, endDate, exportWordPath):
    # Получаем список судей
    judgesList = GetJudgesData()
    
    # Считаем статистику для выгружаемого списка
    statisticsObject = CalculateStatistic(events)
    print(statisticsObject)

    dateDiapazomeBagin = datetime.fromtimestamp(beginDate / 1000)
    dateDiapazomeEnd = datetime.fromtimestamp(endDate / 1000)

    # print("ivent object: " + str(events[0]))
    # print("ID: " + str(events[0][0]))

    doc = Document()  
    # Добавить заголовок  
    heading = doc.add_heading('ГРАФИК ПРОВЕДЕНИЯ ВИДЕОКОНФЕРЕНЦ-СВЯЗИ В ХОЛМСКОМ ГОРОДСКОМ СУДЕ', 0)
    heading.style = doc.styles['Heading 1']
    # Добавить абзац  
    doc.add_paragraph(f'''Выборка на диапазон дат: {dateDiapazomeBagin} - {dateDiapazomeEnd}''')
    doc.add_paragraph("")
    headingStatistic = doc.add_heading('Статистика для данной таблицы:', 0)
    headingStatistic.style = doc.styles['Heading 2']
    doc.add_paragraph(f'''Общее количество ВКС: {statisticsObject['totalEvents']}''')
    doc.add_paragraph(f'''Исходящих ВКС: {statisticsObject['totalOutputsCount']}''')
    doc.add_paragraph(f'''Входящих ВКС: {statisticsObject['totalInputsCount']}''')
    doc.add_paragraph(f'''ВКС во 2 зале: {statisticsObject['totalHall2Count']}''')
    doc.add_paragraph(f'''ВКС в 4 зале: {statisticsObject['totalHall4Count']}''')

    doc.add_paragraph("")
    headingStatistic = doc.add_heading('Таблица ВКС:', 0)
    headingStatistic.style = doc.styles['Heading 2']

    table = doc.add_table(rows=1, cols=5)
    table.style = doc.styles['Table Grid']
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Дата и время начала'
    hdr_cells[1].text = '№ Зала'
    hdr_cells[2].text = 'Ф.И.О.судьи'
    hdr_cells[3].text = 'Суд, рассматривающий дело (учреждение, УФСИН)'
    hdr_cells[4].text = 'Сущность заявки (предмет, допрос свидетеля по делу, ходатайство в порядке ст. 396-399 УПК РФ и т.д.)'

    for event in events:
        dateIvent = ''
        if (event[1]):
            converted_date = datetime.fromtimestamp(int(event[1]) / 1000)
            # let iventDateTime = new Date( parseInt(ivent[1]) ).toISOString()
            # iventDate = new Date(iventDateTime).toLocaleDateString()
            # iventTime = new Date(iventDateTime).toLocaleTimeString([], { hour: "2-digit", minute: "2-digit" })
            dateIvent = str(converted_date)
        else: dateIvent = 'not date'
        
        row_cells = table.add_row().cells
        row_cells[0].text = dateIvent
        row_cells[1].text = event[5]
        row_cells[2].text = str(FindJudgeFromList(judgesList, int(event[4])))
        row_cells[3].text = event[2]
        row_cells[4].text = event[6]

    doc.add_page_break()

    docName = str(GenerateWordName()) # Генерируем имя файла
    print(f'Выгрузка в Word файл ВКС событий, по диапазону дат - {docName}')
    # Сохранить документ
    doc.save(f'{exportWordPath}\{docName}')

    return True

def FindJudgeFromList(list, id):
    for elem in list:
        if elem['Id'] == id:
            return elem['name']

def GenerateWordName():
    generateDate = datetime.now()
    year = generateDate.year
    month = generateDate.month
    day = generateDate.day
    hour = generateDate.hour
    minute = generateDate.minute
    second =generateDate.second
    return f'{day}-{month}-{year}_{hour}-{minute}-{second}.docx'

def CalculateStatistic(eventsArray):
    statisticsObj = {
        "totalEvents": len(eventsArray),
        "totalOutputsCount": 0,
        "totalInputsCount": 0,
        "totalHall2Count": 0,
        "totalHall4Count": 0
    }
    for elem in eventsArray:
        print(elem[2])
        if elem[3] == 'Outbox': statisticsObj['totalOutputsCount'] = statisticsObj['totalOutputsCount']+1
        if elem[3] == 'Inbox': statisticsObj['totalInputsCount'] = statisticsObj['totalInputsCount']+1
        if elem[5] == '2': statisticsObj['totalHall2Count'] = statisticsObj['totalHall2Count']+1
        if elem[5] == '4': statisticsObj['totalHall4Count'] = statisticsObj['totalHall4Count']+1

    return statisticsObj